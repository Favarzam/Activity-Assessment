"""Build a comprehensive .docx findings report for the SIOP second-rater app.

The report bundles (in this order):

1. Title + metadata (generated-at, input files, activities / features counts).
2. Executive summary — pooled reliability metrics (if human ratings available).
3. Mean score per feature — algorithm, human, and delta.
4. Per-feature reliability table + 5×5 confusion matrices (if available).
5. Per-activity findings — for every activity, the text plus a per-feature
   table containing the 0–4 score, human score (if provided), delta, plain-
   English reasoning, and the actual text excerpts that drove the decision.

Only one public function, ``build_docx_report``, which returns the .docx bytes.
"""

from __future__ import annotations

import io
import sys
from datetime import datetime
from pathlib import Path
from typing import Any, Iterable

import pandas as pd

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

_HERE = Path(__file__).resolve().parent
if str(_HERE) not in sys.path:
    sys.path.insert(0, str(_HERE))

from codebook import FEATURE_NAMES, extract_evidence, rationale


# ---------------------------------------------------------------------------
# Small docx helpers
# ---------------------------------------------------------------------------


def _set_default_font(doc: Document, name: str = "Calibri", size_pt: int = 11) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)


def _set_column_widths(table, widths_inches: list[float]) -> None:
    """Best-effort column width setter (Word may still auto-resize)."""
    from docx.shared import Inches

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_inches):
                cell.width = Inches(widths_inches[idx])


def _add_header_row(table, labels: Iterable[str]) -> None:
    hdr = table.rows[0].cells
    for i, label in enumerate(labels):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(str(label))
        run.bold = True
        run.font.size = Pt(10)


def _set_cell_text(cell, text: str, *, bold: bool = False, size_pt: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("" if text is None else str(text))
    run.bold = bold
    run.font.size = Pt(size_pt)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def _fmt(val: Any) -> str:
    if val is None:
        return ""
    if isinstance(val, float):
        if pd.isna(val):
            return ""
        return f"{val:.3f}" if abs(val) < 100 else f"{val:.1f}"
    if pd.isna(val) if not isinstance(val, (list, dict, tuple)) else False:
        return ""
    return str(val)


def _add_df_table(
    doc: Document,
    df: pd.DataFrame,
    *,
    widths_inches: list[float] | None = None,
    max_rows: int | None = None,
    style: str = "Light Grid Accent 1",
) -> None:
    """Render a pandas DataFrame as a simple docx table."""
    if df is None or df.empty:
        doc.add_paragraph("— no data —")
        return
    work = df.reset_index() if df.index.name else df.copy()
    if max_rows is not None:
        work = work.head(max_rows)
    cols = list(work.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    try:
        table.style = style
    except KeyError:
        pass
    _add_header_row(table, cols)
    for _, row in work.iterrows():
        cells = table.add_row().cells
        for i, col in enumerate(cols):
            _set_cell_text(cells[i], _fmt(row[col]))
    if widths_inches:
        _set_column_widths(table, widths_inches)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------


def _section_title_page(
    doc: Document,
    meta: dict[str, Any],
) -> None:
    title = doc.add_heading("SIOP Deterministic Second-Rater", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Findings Report")
    run.italic = True
    run.font.size = Pt(14)

    ts = meta.get("generated_at") or datetime.now().strftime("%Y-%m-%d %H:%M")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Generated: {ts}").font.size = Pt(10)

    doc.add_paragraph()

    info_pairs = [
        ("Activities scored", meta.get("n_activities", "—")),
        ("Features scored", meta.get("n_features", "—")),
        ("Cells produced", meta.get("n_cells", "—")),
        ("Human ratings provided", "yes" if meta.get("has_human") else "no"),
        ("Activity texts source", meta.get("texts_source") or "—"),
        ("Human ratings source", meta.get("human_source") or "—"),
        ("Excluded features included in reliability",
         "yes" if meta.get("include_excluded") else "no"),
    ]
    table = doc.add_table(rows=0, cols=2)
    try:
        table.style = "Light List Accent 1"
    except KeyError:
        pass
    for k, v in info_pairs:
        row = table.add_row().cells
        _set_cell_text(row[0], k, bold=True)
        _set_cell_text(row[1], _fmt(v))
    _set_column_widths(table, [2.5, 3.5])


def _section_executive_summary(
    doc: Document,
    pooled: dict[str, float] | None,
    overall_algo_mean: float,
) -> None:
    doc.add_heading("Executive summary", level=1)
    bullet = doc.add_paragraph(style="List Bullet")
    bullet.add_run(
        f"Mean algorithm score across all activities × features: "
    )
    bullet.add_run(f"{overall_algo_mean:.2f}").bold = True

    if pooled is not None and len(pooled) > 0:
        for key, label in [
            ("AC2_ordinal", "Pooled Gwet's AC2 (ordinal)"),
            ("kappa_w", "Pooled weighted Cohen's κ"),
            ("alpha_ordinal", "Pooled Krippendorff's α (ordinal)"),
            ("exact_pct", "Pooled exact-agreement %"),
            ("within_1_pct", "Pooled within-1 agreement %"),
        ]:
            if key in pooled and pooled[key] is not None:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{label}: ")
                val = pooled[key]
                text = f"{val:.3f}" if key not in ("exact_pct", "within_1_pct") else f"{val:.1f}%"
                p.add_run(text).bold = True
    else:
        doc.add_paragraph(
            "Reliability metrics were not computed because no human ratings "
            "were provided.",
            style="Intense Quote",
        )


def _section_means_per_feature(
    doc: Document,
    algo_wide: pd.DataFrame,
    human_wide: pd.DataFrame | None,
) -> None:
    doc.add_heading("Mean score per feature", level=1)
    means = algo_wide.mean().round(2).to_frame("Algorithm mean")
    means.insert(0, "Feature name",
                 [FEATURE_NAMES[int(c[1:])] for c in means.index])
    means.insert(0, "Feature", means.index.tolist())
    if human_wide is not None:
        h_means = human_wide.mean().round(2).to_frame("Human mean")
        means = means.join(h_means, how="left")
        means["Δ (algo − human)"] = (
            means["Algorithm mean"] - means["Human mean"]
        ).round(2)
    means = means.reset_index(drop=True)
    _add_df_table(doc, means, widths_inches=[0.6, 2.6, 1.0, 1.0, 1.2])


def _section_reliability(
    doc: Document,
    reliability_report: pd.DataFrame,
    algo_wide: pd.DataFrame,
    human_wide: pd.DataFrame,
) -> None:
    doc.add_heading("Reliability (algorithm vs. human)", level=1)
    doc.add_paragraph(
        "Per-feature inter-rater reliability. AC2_ordinal is Gwet's AC2 with "
        "ordinal weights (95 % bootstrap CI). kappa_w is quadratic-weighted "
        "Cohen's κ. alpha_ordinal is Krippendorff's α (ordinal). SMD is "
        "(algorithm − human) standardized mean difference."
    )
    _add_df_table(doc, reliability_report.copy())

    doc.add_heading("Confusion matrices per feature (5 × 5)", level=2)
    doc.add_paragraph(
        "Rows = human score, columns = algorithm score. Both axes span 0–4."
    )
    common_idx = algo_wide.index.intersection(human_wide.index)
    common_cols = [c for c in algo_wide.columns if c in human_wide.columns]
    reported_features = set(reliability_report["feature"].tolist())
    for col in sorted(common_cols, key=lambda c: int(c[1:])):
        if col not in reported_features:
            continue
        fnum = int(col[1:])
        h = human_wide[col].dropna().astype(int)
        a = algo_wide[col].dropna().astype(int)
        idx = h.index.intersection(a.index)
        if len(idx) == 0:
            continue
        h, a = h.loc[idx], a.loc[idx]
        cm = pd.crosstab(h, a, rownames=["Human"], colnames=["Algorithm"])
        for i in range(5):
            if i not in cm.index:
                cm.loc[i] = 0
            if i not in cm.columns:
                cm[i] = 0
        cm = cm.reindex(index=range(5), columns=range(5)).fillna(0).astype(int)
        doc.add_heading(f"{col} — {FEATURE_NAMES.get(fnum, '')}", level=3)
        table = doc.add_table(rows=6, cols=6)
        try:
            table.style = "Light Grid Accent 1"
        except KeyError:
            pass
        _set_cell_text(table.cell(0, 0), "Human \\ Algo", bold=True)
        for j in range(5):
            _set_cell_text(table.cell(0, j + 1), str(j), bold=True)
        for i in range(5):
            _set_cell_text(table.cell(i + 1, 0), str(i), bold=True)
            for j in range(5):
                _set_cell_text(table.cell(i + 1, j + 1), str(int(cm.iloc[i, j])))


def _section_per_activity(
    doc: Document,
    activities: pd.DataFrame,
    algo_wide: pd.DataFrame,
    human_wide: pd.DataFrame | None,
    diags_by_activity: dict[int, dict[int, dict[str, Any]]],
    max_evidence_per_feature: int = 3,
) -> None:
    doc.add_heading("Per-activity findings", level=1)
    doc.add_paragraph(
        "For every activity, the algorithm's 0–4 score for each of the 30 "
        "SIOP features, along with a plain-English reason for the score and "
        "the text excerpts that triggered the indicator. When human ratings "
        "are available, the human score and the algo − human delta are "
        "included."
    )

    idx = activities.set_index("activity_id")
    for aid in sorted(activities["activity_id"].tolist()):
        text = idx.loc[aid, "text"]
        n_chars = int(idx.loc[aid, "characters"]) if "characters" in idx.columns else len(text or "")
        diags = diags_by_activity.get(aid, {})

        doc.add_heading(f"Activity {aid}", level=2)

        meta_p = doc.add_paragraph()
        meta_p.add_run(f"{n_chars:,} characters").italic = True

        doc.add_heading("Activity text", level=3)
        q = doc.add_paragraph(text or "(empty text)", style="Quote")
        for run in q.runs:
            run.font.size = Pt(10)

        doc.add_heading("Per-feature findings", level=3)
        cols = ["Feature", "Name", "Score"]
        has_human = human_wide is not None and aid in human_wide.index
        if has_human:
            cols += ["Human", "Δ"]
        cols += ["Reasoning", "Evidence (text excerpts)"]
        table = doc.add_table(rows=1, cols=len(cols))
        try:
            table.style = "Light Grid Accent 1"
        except KeyError:
            pass
        _add_header_row(table, cols)

        for fnum in range(1, 31):
            feat_col = f"F{fnum}"
            if feat_col not in algo_wide.columns:
                continue
            score = int(algo_wide.loc[aid, feat_col])
            diag = diags.get(fnum, {})
            reason = rationale(fnum, score, diag)

            evidence = extract_evidence(text or "", fnum)
            evidence = evidence[:max_evidence_per_feature]
            if evidence:
                evidence_text = "\n".join(
                    f"• [{e['category']}] {e['snippet']}" for e in evidence
                )
            else:
                evidence_text = "— no matching excerpts —"

            row_vals = [feat_col, FEATURE_NAMES.get(fnum, ""), str(score)]
            if has_human:
                hv = human_wide.loc[aid, feat_col]
                if pd.notna(hv):
                    row_vals += [str(int(hv)), str(int(score - int(hv)))]
                else:
                    row_vals += ["—", "—"]
            row_vals += [reason, evidence_text]

            cells = table.add_row().cells
            for i, val in enumerate(row_vals):
                _set_cell_text(cells[i], val, size_pt=9)

        widths = [0.55, 1.5, 0.45]
        if has_human:
            widths += [0.45, 0.35]
        widths += [1.7, 2.1]
        _set_column_widths(table, widths)


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------


def build_docx_report(
    *,
    activities: pd.DataFrame,
    algo_wide: pd.DataFrame,
    diags_by_activity: dict[int, dict[int, dict[str, Any]]],
    human_wide: pd.DataFrame | None = None,
    reliability_report: pd.DataFrame | None = None,
    pooled: dict[str, float] | None = None,
    meta: dict[str, Any] | None = None,
    max_evidence_per_feature: int = 3,
) -> bytes:
    """Assemble the full findings report and return .docx bytes.

    Parameters
    ----------
    activities
        DataFrame with ``activity_id``, ``text`` (and optionally
        ``characters``) columns — one row per activity.
    algo_wide
        Wide-form algorithm scores: index = ``activity_id``, columns = ``F1``
        … ``F30``.
    diags_by_activity
        ``{activity_id: {feature_num: diag_dict, …}, …}``. Typically built by
        calling ``codebook.score_activity(text)`` per activity and keeping
        the second element of the returned tuple.
    human_wide
        Optional human ratings, same shape as ``algo_wide``.
    reliability_report, pooled
        Output of ``reliability.per_feature_report`` and
        ``reliability.pooled_summary``.
    meta
        Free-form metadata dict for the title page.
    max_evidence_per_feature
        Cap on the number of text excerpts shown per feature per activity.
    """
    doc = Document()
    _set_default_font(doc)
    meta = meta or {}

    if pooled is not None and hasattr(pooled, "to_dict"):
        pooled = {k: (None if pd.isna(v) else float(v))
                  for k, v in pooled.to_dict().items()}
    elif pooled is None:
        pooled = None
    else:
        pooled = dict(pooled)

    overall_algo_mean = float(algo_wide.mean().mean())
    meta = {
        **meta,
        "n_activities": len(activities),
        "n_features": len(algo_wide.columns),
        "n_cells": len(activities) * len(algo_wide.columns),
        "has_human": human_wide is not None,
        "generated_at": meta.get("generated_at")
        or datetime.now().strftime("%Y-%m-%d %H:%M"),
    }

    _section_title_page(doc, meta)

    doc.add_page_break()
    _section_executive_summary(doc, pooled, overall_algo_mean)

    doc.add_page_break()
    _section_means_per_feature(doc, algo_wide, human_wide)

    if reliability_report is not None and not reliability_report.empty and human_wide is not None:
        doc.add_page_break()
        _section_reliability(doc, reliability_report, algo_wide, human_wide)

    doc.add_page_break()
    _section_per_activity(
        doc,
        activities,
        algo_wide,
        human_wide,
        diags_by_activity,
        max_evidence_per_feature=max_evidence_per_feature,
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
